#!/usr/bin/env python3
"""
Generate the 3DShoemaker Rhino 8 Plugin User Manual as a Word document.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_manual():
    doc = Document()

    # -- Page setup --
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # -- Custom styles --
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    # =====================================================================
    # COVER PAGE
    # =====================================================================
    for _ in range(6):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('3DShoemaker')
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Rhino 8 Plugin - User Manual')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x4A, 0x6C, 0x9E)

    doc.add_paragraph('')

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('Version 8.4.0.8')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')

    url = doc.add_paragraph()
    url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = url.add_run('https://ShoeLastMaker.com')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x33, 0x66, 0xCC)

    doc.add_page_break()

    # =====================================================================
    # TABLE OF CONTENTS
    # =====================================================================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('1.', 'Introduction'),
        ('2.', 'System Requirements'),
        ('3.', 'Installation'),
        ('4.', 'Licensing & Activation'),
        ('5.', 'Getting Started'),
        ('6.', 'Last Commands'),
        ('7.', 'Morph Commands'),
        ('8.', 'Component Commands'),
        ('9.', 'Grading Commands'),
        ('10.', 'Foot Analysis Commands'),
        ('11.', 'Orthotic Commands'),
        ('12.', 'Sandal Commands'),
        ('13.', 'Editing Commands'),
        ('14.', 'Parameterization Commands'),
        ('15.', 'View & Display Commands'),
        ('16.', 'Export & Utility Commands'),
        ('17.', 'Data Models & Parameters'),
        ('18.', 'Layer Structure'),
        ('19.', 'Document Persistence'),
        ('20.', 'Troubleshooting'),
        ('21.', 'Terms and Conditions'),
    ]
    for num, item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}  {item}')
        run.font.size = Pt(12)

    doc.add_page_break()

    # =====================================================================
    # 1. INTRODUCTION
    # =====================================================================
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        '3DShoemaker is a comprehensive footwear design plugin for Rhinoceros 3D (Rhino 8). '
        'It provides a complete suite of tools for designing shoe lasts, insoles, soles, heels, '
        'orthotics, sandals, and other footwear components. The plugin integrates directly into '
        'the Rhino 8 command line and provides over 90 specialized commands for footwear CAD.'
    )
    doc.add_paragraph(
        'This Python 3 edition of 3DShoemaker is designed for Rhino 8 and leverages the '
        'RhinoCommon API, rhinoscriptsyntax, and Eto.Forms for cross-platform compatibility '
        'on both Windows and macOS.'
    )

    doc.add_heading('Key Features', level=2)
    features = [
        'Shoe last creation, parameterization, and modification',
        'Advanced morphing (FFD, Point-to-Point, Mesh-to-Mesh, Twist)',
        'Component creation: insoles, soles, heels, shank boards, met pads, shoe trees',
        'Size grading across EU, US, UK, and Mondopoint systems',
        'Foot scan import and plantar analysis',
        'Orthotic design with arch support, heel cupping, and twist adjustments',
        'Sandal construction with groove and thong slot features',
        'Interactive surface sculpting and curve editing',
        'Comprehensive parameterization for fit customization',
        'Print preparation for 3D printing and vacuum forming',
        'Alpha joint and rail guide joint assembly systems',
        'Batch grading to multiple sizes simultaneously',
        'JSON parameter import/export for workflow automation',
        'Cryptolens license management (Personal, Business, Enterprise editions)',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_page_break()

    # =====================================================================
    # 2. SYSTEM REQUIREMENTS
    # =====================================================================
    doc.add_heading('2. System Requirements', level=1)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    cells = table.rows[0].cells
    cells[0].text = 'Requirement'
    cells[1].text = 'Details'
    for cell in cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    data = [
        ('Software', 'Rhinoceros 3D Version 8.0 or later'),
        ('Operating System', 'Windows 10/11 (64-bit) or macOS 12+'),
        ('Python', 'Python 3 (included with Rhino 8)'),
        ('Internet', 'Required for license activation and validation'),
    ]
    for i, (key, val) in enumerate(data):
        cells = table.rows[i + 1].cells
        cells[0].text = key
        cells[1].text = val

    doc.add_page_break()

    # =====================================================================
    # 3. INSTALLATION
    # =====================================================================
    doc.add_heading('3. Installation', level=1)

    doc.add_heading('Automatic Installation', level=2)
    doc.add_paragraph(
        'The plugin includes an installer script that automatically detects your Rhino 8 '
        'installation and copies the plugin files to the correct location.'
    )

    doc.add_heading('Windows', level=3)
    p = doc.add_paragraph()
    run = p.add_run('python install.py')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    doc.add_paragraph(
        'Files are installed to: %APPDATA%\\McNeel\\Rhinoceros\\8.0\\Plug-ins\\PythonPlugIns\\3DShoemaker\\'
    )

    doc.add_heading('macOS', level=3)
    p = doc.add_paragraph()
    run = p.add_run('python3 install.py')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    doc.add_paragraph(
        'Files are installed to: ~/Library/Application Support/McNeel/Rhinoceros/8.0/Plug-ins/PythonPlugIns/3DShoemaker/'
    )

    doc.add_heading('Uninstallation', level=2)
    p = doc.add_paragraph()
    run = p.add_run('python install.py --uninstall')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    doc.add_paragraph(
        'This removes all plugin files from the Rhino 8 plugin directory. '
        'Restart Rhino 8 after installing or uninstalling.'
    )

    doc.add_heading('Custom Destination', level=2)
    p = doc.add_paragraph()
    run = p.add_run('python install.py --dest /path/to/custom/directory')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    doc.add_paragraph('Use the --dest flag to specify a custom installation directory for CI/testing.')

    doc.add_heading('Installed File Structure', level=2)
    structure_lines = [
        '3DShoemaker/',
        '  __init__.py              Plugin entry point',
        '  manifest.yml             Plugin metadata',
        '  Terms.txt                License agreement',
        '  plugin/',
        '    __init__.py            Package init',
        '    plugin_main.py         Main plugin singleton',
        '    document_settings.py   Document settings manager',
        '    material_thicknesses.py Material parameters',
        '    preview_module.py      Custom display conduit',
        '    commands/              All command modules (12 files)',
        '    models/                Data models (Last, Insert, Bottom, Foot)',
        '    forms/                 Eto.Forms UI dialogs (12 files)',
        '    utils/                 Utility modules (7 files)',
    ]
    for line in structure_lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

    doc.add_page_break()

    # =====================================================================
    # 4. LICENSING & ACTIVATION
    # =====================================================================
    doc.add_heading('4. Licensing & Activation', level=1)
    doc.add_paragraph(
        '3DShoemaker uses the Cryptolens licensing system. Licenses are available in three editions:'
    )

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    cells = table.rows[0].cells
    cells[0].text = 'Edition'
    cells[1].text = 'Usage'
    cells[2].text = 'Features'
    for cell in cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    editions = [
        ('Personal', 'Non-commercial use only', 'All design tools, single machine'),
        ('Business', 'Commercial/professional use', 'All design tools, email support, single machine'),
        ('Enterprise', 'Commercial with extended support', 'All tools, priority support, training options'),
    ]
    for i, (ed, usage, feat) in enumerate(editions):
        cells = table.rows[i + 1].cells
        cells[0].text = ed
        cells[1].text = usage
        cells[2].text = feat

    doc.add_heading('Activating Your License', level=2)
    doc.add_paragraph('1. Launch Rhino 8. The plugin loads automatically.')
    doc.add_paragraph('2. Type Activate3DShoemaker in the Rhino command line.')
    doc.add_paragraph('3. Enter your license key when prompted.')
    doc.add_paragraph('4. The plugin contacts the license server to validate your key.')
    doc.add_paragraph('5. On success, the license is cached locally for offline use.')

    doc.add_heading('Deactivating Your License', level=2)
    doc.add_paragraph(
        'To transfer your license to a different machine, type Deactivate3DShoemaker in the '
        'command line and confirm with "YES". This frees the machine slot on the server.'
    )

    doc.add_heading('Offline Mode', level=2)
    doc.add_paragraph(
        'Once activated, the plugin caches the license locally. If the license server is '
        'unreachable, the plugin falls back to the cached license. An internet connection '
        'is required for the initial activation.'
    )

    doc.add_page_break()

    # =====================================================================
    # 5. GETTING STARTED
    # =====================================================================
    doc.add_heading('5. Getting Started', level=1)

    doc.add_heading('Plugin Initialization', level=2)
    doc.add_paragraph(
        'When Rhino 8 starts, the 3DShoemaker plugin initializes automatically. You will see '
        'messages in the command line confirming the plugin version and the number of registered '
        'commands. The plugin sets up:'
    )
    items = [
        'Layer hierarchy (SLM parent layer with child layers for each object class)',
        'Default rendering settings (neutral studio lighting, white background)',
        'Perspective viewport camera angle optimized for shoe last viewing',
        'Object classification layers (Last, Insert, Bottom, Foot)',
        'Document event hooks for automatic save/load of plugin data',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Basic Workflow', level=2)
    steps = [
        ('Create a New Build', 'Use the NewBuild command to set up a new shoe last with initial parameters (size, heel height, toe shape, etc.).'),
        ('Import or Design', 'Import an existing last file (ImportLast) or use the interactive builder. Import foot scans with ImportFoot.'),
        ('Add Components', 'Create insole (CreateInsole), sole (CreateSole), heel (CreateHeel), and other components.'),
        ('Adjust Parameters', 'Fine-tune fit with AdjustFitCustomization, material thicknesses, and cross-section planes.'),
        ('Morph & Edit', 'Use Morph for shape transformation, Sculpt for freeform editing, or EditCurve for curve refinement.'),
        ('Grade Sizes', 'Grade to additional sizes with GradeFootwear or BatchGrade.'),
        ('Export & Produce', 'Export parameters, prepare for 3D printing (PrintPrep), or create vacuum form shells.'),
    ]
    for title_text, desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(f'{title_text}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # =====================================================================
    # 6. LAST COMMANDS
    # =====================================================================
    doc.add_heading('6. Last Commands', level=1)
    doc.add_paragraph(
        'Last commands handle the creation, modification, import, export, and parameterization '
        'of shoe lasts - the core 3D form around which all footwear is designed.'
    )

    last_commands = [
        ('NewBuild', 'Creates a new shoe last build from parameters. Opens an interactive dialog '
         'where you can specify shoe size, size system (EU/US/UK/Mondopoint), heel height, toe shape, '
         'ball width, instep height, and other fundamental measurements. The last is generated as a '
         'NURBS surface and placed on the SLM::Last layer.'),
        ('NewBuildScriptable', 'Non-interactive version of NewBuild for scripting and automation. '
         'Accepts all parameters via command-line options, enabling batch processing and integration '
         'with external scripts.'),
        ('UpdateLast', 'Updates existing last geometry after parameter changes. Regenerates the last '
         'surface while preserving its position and associations with other components.'),
        ('ImportLast', 'Imports a last from an external file. Supports 3DM, STEP, IGES, STL, OBJ, '
         'and other common CAD formats. The imported geometry is placed on the Last layer and registered '
         'with the plugin.'),
        ('ExportLast', 'Exports the current last to a file. Supports 3DM, STEP, IGES, STL, and OBJ '
         'formats. Includes an option to export with or without component geometry.'),
        ('GradeLast', 'Grades (scales) the last to a different size. Uses the configured size system '
         'and applies proportional scaling with girth adjustments.'),
        ('FlattenLast', 'Flattens the last bottom to a 2D pattern. Creates an unrolled representation '
         'of the bottom surface suitable for pattern cutting.'),
        ('GazeAtLast', 'Sets the viewport to standard last-viewing angles. Cycles through predefined '
         'camera positions: front, back, medial, lateral, top, and bottom views.'),
        ('ChangeLastParameterization', 'Opens a dialog to modify last parameters including stick length, '
         'ball width, heel width, toe spring, heel height, ball break angle, and girth measurements.'),
        ('ExportLastParameters', 'Exports all last parameters to a JSON file for documentation, '
         'sharing, or use in external tools.'),
        ('ImportParameters', 'Imports parameters from a JSON file and applies them to the current last.'),
        ('ExportMeasurementEquations', 'Exports the measurement equations that define relationships '
         'between parameters (e.g., how ball girth relates to size).'),
        ('NameObjectsInDoc', 'Assigns standardized names to all objects in the document based on their '
         'layer and function. Ensures consistency for scripting and automation.'),
        ('GetObjectIDName', 'Reports the internal name and GUID of a selected object.'),
        ('Establish', 'Initializes a new shoe last project. Sets up the complete layer hierarchy, '
         'default settings, and prepares the document for footwear design.'),
    ]

    for cmd_name, desc in last_commands:
        doc.add_heading(cmd_name, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # =====================================================================
    # 7. MORPH COMMANDS
    # =====================================================================
    doc.add_heading('7. Morph Commands', level=1)
    doc.add_paragraph(
        'Morph commands provide shape transformation tools for modifying last and component '
        'geometry. These are essential for customizing fits, adapting designs to different '
        'foot shapes, and making localized adjustments.'
    )

    doc.add_heading('Morph', level=2)
    doc.add_paragraph(
        'The main morph command offers seven morphing operations accessible via command-line options:'
    )

    morph_ops = [
        ('FFD (Free-Form Deformation)', 'Select objects, pick source control points, then pick '
         'matching target points. The geometry deforms smoothly to map source positions to targets '
         'using weighted inverse-distance interpolation.'),
        ('PointToPoint (P2P)', 'Morph a mesh using explicit point correspondences. Each vertex is '
         'displaced based on the weighted contribution of all source-to-target vectors.'),
        ('MeshToMesh (M2M)', 'Morph by establishing closest-point correspondences between a source '
         'reference mesh and a target reference mesh. Ideal for adapting a last to match a foot scan.'),
        ('MeshToPoint (M2P)', 'Hybrid approach using a source mesh and target pick-points.'),
        ('ForeFootTwist', 'Applies twist deformation to the forefoot region using TwistSpaceMorph. '
         'Specify the axis origin, twist angle (degrees), and the distance range over which the twist ramps.'),
        ('RearFootTwist', 'Similar to ForeFootTwist but oriented toward the heel-to-midfoot section. '
         'Used for rearfoot posting adjustments.'),
        ('NurbsSurface', 'Point-to-point morph specifically for NURBS surfaces, preserving surface continuity.'),
    ]
    for op_name, desc in morph_ops:
        p = doc.add_paragraph()
        run = p.add_run(f'{op_name}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('NewMorph', level=2)
    doc.add_paragraph(
        'Opens an interactive Eto dialog for visual morphing. The dialog presents a dropdown '
        'for selecting the morph operation, a tolerance control, and a twist angle input. '
        'After clicking "Apply Morph", the selected operation executes with the specified parameters.'
    )

    doc.add_heading('NewMorphScriptable', level=2)
    doc.add_paragraph(
        'Scriptable (non-interactive) morph command. Accepts the morph type as a string '
        '(P2P, M2M, M2P, FFD, ForeFootTwist, RearFootTwist, NURBS) and a Tolerance option. '
        'Designed for automation and batch processing.'
    )

    doc.add_page_break()

    # =====================================================================
    # 8. COMPONENT COMMANDS
    # =====================================================================
    doc.add_heading('8. Component Commands', level=1)
    doc.add_paragraph(
        'Component commands create the individual parts that make up a complete shoe. Each component '
        'is generated on its own layer within the SLM hierarchy.'
    )

    component_commands = [
        ('CreateInsole', 'Generates insole geometry from the last bottom surface. Uses the current '
         'material thickness settings and parameterization to create a properly contoured insole board.'),
        ('CreateSole', 'Creates the outsole geometry. Options include sole thickness, profile type '
         '(flat, rocker, wedge), and tread pattern integration.'),
        ('CreateHeel', 'Generates heel geometry based on the heel height, seat shape, and breast '
         'line position. The heel is positioned relative to the last bottom.'),
        ('CreateHeelParts', 'Creates individual heel sub-components including top lift, heel stack, '
         'and heel breast. Each part can be independently adjusted.'),
        ('CreateTopPiece', 'Creates the top piece (top lift) - the bottom-most layer of the heel '
         'that contacts the ground.'),
        ('CreateShankBoard', 'Generates the shank board that provides structural support between '
         'the heel and ball of the shoe. Thickness is configurable via material settings.'),
        ('CreateMetPad', 'Creates a metatarsal pad positioned behind the metatarsal heads. '
         'Parameters include size, height, and position offset.'),
        ('CreateToeCrest', 'Generates a toe crest feature for orthotic or comfort applications.'),
        ('CreateToeRidge', 'Creates a toe ridge element along the toe line of the last.'),
        ('CreateThongHole', 'Creates a thong hole in the sole for sandal construction.'),
        ('CreatePinHole', 'Creates precision pin holes for manufacturing alignment.'),
        ('CreateShoeTree', 'Generates a shoe tree form based on the current last shape. '
         'The shoe tree is slightly smaller than the last interior volume.'),
        ('CreateUpperBodies', 'Creates upper body component surfaces based on the last shape '
         'and design lines. Used for pattern development.'),
        ('MakeComponent', 'Generic component creation command that allows selecting from all '
         'available component types via a dropdown menu.'),
        ('CreateAlphaJoint', 'Creates an alpha joint assembly - a hinge mechanism for jointed '
         'shoe lasts that allows the forefoot and rearfoot sections to articulate.'),
        ('CreateRailGuideJoint', 'Creates a rail guide joint assembly - a sliding mechanism '
         'that constrains motion along a defined path for adjustable last designs.'),
        ('CreateMockup', 'Creates a 3D mockup combining all current components into a unified '
         'visualization. Useful for design review and presentation.'),
    ]

    for cmd_name, desc in component_commands:
        doc.add_heading(cmd_name, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # =====================================================================
    # 9. GRADING COMMANDS
    # =====================================================================
    doc.add_heading('9. Grading Commands', level=1)
    doc.add_paragraph(
        'Grading commands scale footwear designs between sizes. The grading system supports '
        'multiple size systems and maintains proportional relationships between dimensions.'
    )

    doc.add_heading('Size Systems', level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'
    cells = table.rows[0].cells
    cells[0].text = 'System'
    cells[1].text = 'Base Size'
    cells[2].text = 'Base Length (mm)'
    cells[3].text = 'Increment (mm/size)'
    for cell in cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    size_data = [
        ('EU', '40', '260.0', '6.667'),
        ('US', '8', '260.0', '8.467'),
        ('UK', '7', '260.0', '8.467'),
        ('Mondopoint', '260', '260.0', '5.0'),
    ]
    for i, row_data in enumerate(size_data):
        cells = table.rows[i + 1].cells
        for j, val in enumerate(row_data):
            cells[j].text = val

    doc.add_heading('GradeFootwear', level=2)
    doc.add_paragraph(
        'Grades complete footwear to a different size. The command prompts for a target size '
        'and size system, then computes the scale factor and applies it to all components. '
        'Grading affects:'
    )
    grade_items = [
        'Insole geometry (scaled from heel center origin)',
        'Outline/last-outline curves',
        'Third-party insole geometry if present',
        'All objects on SLM class layers',
        'Girth measurements (CBG ball girth and CIG instep girth)',
    ]
    for item in grade_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('BatchGrade', level=2)
    doc.add_paragraph(
        'Batch-grade footwear to multiple target sizes in one operation. Enter a comma-separated '
        'list of target sizes (e.g., "38,39,40,41,42"). For each size, the command duplicates all '
        'SLM-layer objects, applies the grade scale factor, and offsets the copies laterally by '
        '350mm so graded sizes do not overlap. Each duplicate is named with a "_SizeXX" suffix.'
    )

    doc.add_page_break()

    # =====================================================================
    # 10. FOOT ANALYSIS COMMANDS
    # =====================================================================
    doc.add_heading('10. Foot Analysis Commands', level=1)
    doc.add_paragraph(
        'Foot analysis commands handle the import and measurement of foot scan data, '
        'providing the foundation for custom-fit footwear design.'
    )

    doc.add_heading('ImportFoot', level=2)
    doc.add_paragraph(
        'Opens a file dialog for importing foot scan files. Supported formats include:'
    )

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Light Grid Accent 1'
    cells = table.rows[0].cells
    cells[0].text = 'Category'
    cells[1].text = 'Formats'
    for cell in cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    table.rows[1].cells[0].text = 'Mesh Files'
    table.rows[1].cells[1].text = '.stl, .obj, .ply, .3mf, .off, .3dm'
    table.rows[2].cells[0].text = 'Point Clouds'
    table.rows[2].cells[1].text = '.xyz, .pts, .csv, .txt, .asc'

    doc.add_paragraph(
        'Imported geometry is placed on the SLM::Foot layer and the file path is stored '
        'in the document settings.'
    )

    doc.add_heading('OpenImportFootForm', level=2)
    doc.add_paragraph(
        'Alternative entry point to the foot import UI. Delegates to ImportFoot. Provided '
        'as a separate command name for toolbar and alias compatibility.'
    )

    doc.add_heading('AnalyzePlantarFootScan', level=2)
    doc.add_paragraph(
        'Analyzes a selected foot scan mesh and computes plantar metrics:'
    )
    metrics = [
        'Foot Length - Total length from heel to toe',
        'Ball Width - Width at the metatarsal heads (~72% from heel)',
        'Heel Width - Width at ~10% from the rear of the foot',
        'Arch Height - Height of the medial arch at ~60% from heel',
        'Foot Height - Overall vertical extent',
    ]
    for m in metrics:
        doc.add_paragraph(m, style='List Bullet')
    doc.add_paragraph(
        'Results are displayed in the command line and stored in document settings. Visual '
        'measurement annotation lines are created on the SLM::Measurements layer.'
    )

    doc.add_page_break()

    # =====================================================================
    # 11. ORTHOTIC COMMANDS
    # =====================================================================
    doc.add_heading('11. Orthotic Commands', level=1)
    doc.add_paragraph(
        'Orthotic commands provide tools for designing custom orthotic devices, including '
        'full orthotics, adjustments, and 3D printing preparation.'
    )

    orthotic_commands = [
        ('MakeOrthotic', 'Creates an orthotic device from foot/last data. Prompts for arch height, '
         'heel cup depth, posting angles, trim line position, and material thickness. The orthotic '
         'is generated as a contoured surface based on the foot scan data and stored on the SLM::Orthotic layer.'),
        ('AdjustOrthoticToBlank', 'Adjusts the orthotic design to fit within the boundaries of a '
         'manufacturing blank. Prompts for the blank outline curve and trims/adjusts the orthotic '
         'to fit within it.'),
        ('AdjustOrthoticArchHeightAndLength', 'Modifies the arch height and arch length of an existing '
         'orthotic. Allows fine-tuning the medial longitudinal arch support profile.'),
        ('AdjustOrthoticFeature', 'Adjusts a specific feature on the orthotic (met pad, heel post, '
         'forefoot post, lateral wedge, or medial wedge). Select the feature type and enter new dimensions.'),
        ('TwistOrthotic', 'Applies a twist deformation to the orthotic using TwistSpaceMorph. '
         'Useful for adding forefoot or rearfoot posting via deformation rather than material addition.'),
        ('PrintPrepOrthotic', 'Prepares a single orthotic for 3D printing. Converts to mesh, '
         'adds shell thickness, closes holes, and optimizes orientation for the print bed.'),
        ('PrintPrepOrthotics', 'Batch print preparation for multiple orthotics. Processes all '
         'objects on the Orthotic layer and arranges them on the print bed.'),
    ]

    for cmd_name, desc in orthotic_commands:
        doc.add_heading(cmd_name, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # =====================================================================
    # 12. SANDAL COMMANDS
    # =====================================================================
    doc.add_heading('12. Sandal Commands', level=1)
    doc.add_paragraph(
        'Sandal commands provide specialized tools for sandal and open-toe footwear design.'
    )

    sandal_commands = [
        ('BuildSandal', 'Creates a complete sandal from the current last/insert data. Builds the '
         'outsole (extruded from the outline curve) and a contoured footbed with arch support and '
         'heel cupping. Parameters: OutsoleThickness, MidsoleThickness, ArchHeight, HeelCupDepth.'),
        ('BuildInsert', 'Creates a custom insert/footbed for a sandal. Uses the insole outline '
         'to generate a contoured surface with configurable thickness, arch height, heel cup depth, '
         'and top cover. The insert is offset to create solid thickness.'),
        ('AddSandalGroove', 'Adds a groove to a sandal sole for strap attachment. Select the sole, '
         'then select or draw a groove path curve. Configure groove Width and Depth. The groove is '
         'created by sweeping a rectangular cross-section along the path and performing a boolean difference.'),
        ('AddThongSlot', 'Adds a thong slot to a sandal for thong-strap insertion. Pick a position '
         'on the sole (or press Enter for automatic toe-area placement). Configure SlotWidth, SlotDepth, '
         'and SlotLength. Creates a rectangular cut in the sole.'),
        ('ToggleThongSlotInclusion', 'Toggles the visibility of thong slot objects. Shows or hides '
         'all objects with names starting with "ThongSlot".'),
        ('AddMetpad', 'Adds a dome-shaped metatarsal pad to an insert or footbed. Pick the pad center '
         '(or Enter for auto-positioning at ~65% length) and configure Radius and Height. The pad is '
         'boolean-unioned with the insert or added as a separate object.'),
    ]

    for cmd_name, desc in sandal_commands:
        doc.add_heading(cmd_name, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # =====================================================================
    # 13. EDITING COMMANDS
    # =====================================================================
    doc.add_heading('13. Editing Commands', level=1)
    doc.add_paragraph(
        'Editing commands provide geometry modification tools for refining designs.'
    )

    editing_commands = [
        ('EditCurve', 'Enters curve editing mode. Select a curve and its control-point grips '
         'are activated. A red preview conduit shows the curve in real-time. Drag grip points to '
         'modify the curve shape. Use EndEdit to commit or revert changes.'),
        ('EndEdit', 'Ends the current editing session. Choose "Commit" to keep changes or "Revert" '
         'to restore the original geometry. Turns off grips and disables the preview conduit.'),
        ('MoveObjectGrips', 'Moves object control points by a specified vector. Select an object, '
         'select specific grip points (or Enter for all), pick a base point and destination point. '
         'The grips are translated by the resulting vector.'),
        ('Sculpt', 'Interactive sculpting tool for surfaces and meshes. Select an object, configure '
         'BrushRadius, Strength, and Direction (Push/Pull), then click on the surface to deform. '
         'Uses Gaussian falloff for smooth deformation. Works on meshes, BReps, and SubD geometry.'),
        ('BlendSurfaceToSurface', 'Creates a smooth blend surface between two surface edges. '
         'Select an edge on each surface and choose the continuity level (Position, Tangent, or '
         'Curvature/G2). Falls back to lofting if the native blend fails.'),
        ('GirthCurveAveraging', 'Averages multiple girth measurement curves. Select 2 or more '
         'curves, which are rebuilt to the same point count, then control points are averaged to '
         'produce a mean cross-sectional profile. Reports the girth length of the result.'),
        ('AdjustSurfacingCurveControlPointPosition', 'Fine-tune individual control points on a '
         'surfacing curve by numeric entry. Lists all control points with coordinates, prompts for '
         'an index, then allows entering precise X, Y, Z coordinates or picking a new position.'),
        ('CopyObjectToMultiplePoints', 'Copies geometry to multiple locations. Select a source '
         'object, pick a base point, then pick multiple destination points (Enter to finish). '
         'A translated duplicate is placed at each destination.'),
    ]

    for cmd_name, desc in editing_commands:
        doc.add_heading(cmd_name, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # =====================================================================
    # 14. PARAMETERIZATION COMMANDS
    # =====================================================================
    doc.add_heading('14. Parameterization Commands', level=1)
    doc.add_paragraph(
        'Parameterization commands allow adjusting the design parameters that control '
        'geometry generation, material properties, fit, and manufacturing specifications.'
    )

    param_commands = [
        ('ChangeParameter', 'Generic parameter change command. Enter any parameter name and new '
         'value. The value is auto-converted (numbers, booleans, or strings). Triggers a rebuild.'),
        ('ChangeComponentParameterization', 'Adjusts parameters for specific components (Sole, Heel, '
         'ShankBoard, TopPiece, InsoleBoard, Welt, Midsole). Each component has its own set of '
         'adjustable parameters including thickness and profile type.'),
        ('ChangeInsertParameterization', 'Adjusts insert/insole parameters: Thickness, TopCover, '
         'BottomCover, ArchHeight, HeelCupDepth, MedialPosting, LateralPosting, and Material '
         '(EVA, Cork, Leather, Polypropylene, Carbon, Nylon).'),
        ('AdjustBottomComponentParameterization', 'Adjusts bottom component thicknesses: '
         'OutsoleThickness, MidsoleThickness, InsoleBoardThickness, ShankThickness, WeltThickness, '
         'and Profile (Flat, Rocker, Negative). Displays total bottom thickness.'),
        ('AdjustMaterial', 'Adjusts material properties for a component (Insert, Bottom, or Last). '
         'Set the material type, density (g/cm3), and Shore A hardness.'),
        ('AdjustMaterialThicknesses', 'Interactive editor for all material thickness values. Displays '
         'current values and allows editing individual keys or using "all" for a complete editor. '
         'Reports total insole, total bottom, and total build height after changes.'),
        ('AdjustFitCustomization', 'Adjusts fit parameters that control how tightly the footwear '
         'conforms to the foot: ToeRoom, BallEase, InstepEase, HeelEase, WidthEase, and '
         'GirthAdjustment (all in mm).'),
        ('AdjustFootbedDepth', 'Sets the footbed depth (mm) - how deeply the foot sits into the '
         'footbed cavity. Range: 0-30mm.'),
        ('AdjustLastDepthForFootbeds', 'Adjusts the last interior volume to accommodate a footbed. '
         'Manual mode allows specifying the depth adjustment directly. Auto mode calculates from '
         'the total insole thickness. Applies an offset to the last geometry.'),
        ('AdjustCSPlanePositions', 'Repositions cross-section measurement planes along the last. '
         'Standard sections: Heel (0%), Seat (10%), Instep (40%), Waist (50%), Ball (68%), '
         'Toe (90%), TipToe (100%). Values are ratios of total last length from the heel.'),
        ('ChangeStatus', 'Assigns a design status to selected objects: Draft, Review, Approved, '
         'Locked, or Archived. Locked objects are set to locked mode; Archived objects are hidden. '
         'Status is stored in user text on each object.'),
    ]

    for cmd_name, desc in param_commands:
        doc.add_heading(cmd_name, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # =====================================================================
    # 15. VIEW & DISPLAY COMMANDS
    # =====================================================================
    doc.add_heading('15. View & Display Commands', level=1)
    doc.add_paragraph(
        'View commands control visualization, clipping planes, rendering, flattening, '
        'and print preparation.'
    )

    view_commands = [
        ('DrawClippingPlanes', 'Creates clipping planes at cross-section locations on the last '
         '(Ball, Instep, Waist, Waist2, Arch, Arch2, Heel). Each plane is placed on the '
         'SLM::ClippingPlanes layer and named CP_SectionName. If no section data exists, a '
         'default clipping plane is created at the origin.'),
        ('RenderComponents', 'Applies render materials to footwear components based on layer names. '
         'Default colors: Last (tan), Insert (blue), Sole (dark), Heel (charcoal), ShankBoard (brown), '
         'TopPiece (grey), MetPad (light blue). Materials include moderate shine and zero transparency.'),
        ('FlattenInsert', 'Flattens insert geometry to a 2D pattern. Converts the selected surface '
         'to a mesh, projects vertices to Z=0, and offsets the result to the side. Creates outline '
         'curves from naked edges on the SLM::Flattened layer.'),
        ('FlattenSole', 'Flattens sole geometry to a 2D pattern. Uses the same flattening algorithm '
         'as FlattenInsert but applies it to sole surfaces.'),
        ('FlattenBottomSides', 'Flattens multiple bottom side surfaces to 2D patterns. Select '
         'multiple surfaces and each is flattened independently.'),
        ('PrintPrep', 'Prepares model for 3D printing. Opens the PrintPrepForm dialog to configure: '
         'ShellThickness (for hollow shells), MaximizePrintableArea (centers and rotates for optimal '
         'bed usage), and ForPostProcessing (normalizes and unifies mesh normals). Processes selected '
         'BReps, meshes, or SubD objects and outputs print-ready meshes on the SLM::PrintPrep layer.'),
    ]

    for cmd_name, desc in view_commands:
        doc.add_heading(cmd_name, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # =====================================================================
    # 16. EXPORT & UTILITY COMMANDS
    # =====================================================================
    doc.add_heading('16. Export & Utility Commands', level=1)

    export_commands = [
        ('ExportInsertParameters', 'Exports insert parameters to a JSON file. Opens a save dialog '
         'and writes all serializable insert parameters.'),
        ('ExportSupportParameters', 'Exports bottom/support parameters to a JSON file.'),
        ('Open3DShoemakerOptions', 'Opens the 3DShoemaker options/settings dialog (Eto.Forms). '
         'Configure plugin preferences and workspace settings.'),
        ('OpenFolderWatcher', 'Opens a folder watcher window that monitors a directory for new '
         'scan files and automatically imports them when detected.'),
        ('RebuildFootwear', 'Rebuilds all footwear geometry from stored parameters. Regenerates '
         'the last, insert (curves, surfaces, body), and bottom components.'),
        ('VacuumForm', 'Opens the vacuum forming preparation dialog. Select an object, specify '
         'material thickness and draft angle, and create an offset shell suitable for vacuum '
         'forming production.'),
        ('MeasureLast', 'Measures and reports last dimensions including: Length, Ball Width, '
         'Heel Width, Ball/Instep/Waist/Arch/Heel/Ankle Girths, Heel Height, Toe Spring, '
         'Ball Break Angle, Ball Roll Bulge, Ball Line Ratio, and Arch Length. Also measures '
         'actual girth curves in the document.'),
        ('ChangeClippingPlane', 'Modifies the position of an existing clipping plane. Select a '
         'clipping plane and pick a new origin point.'),
        ('SnapCurves', 'Snaps curves to a mesh or surface. Select curves and a target surface; '
         'the curves are projected/snapped to conform to the target geometry.'),
        ('Squeeze', 'Applies squeeze deformation to geometry. Select objects and specify a '
         'squeeze factor (0.5 = compress to half, 2.0 = expand to double).'),
        ('Testing', 'Debug and diagnostic command. Reports plugin status including: whether the '
         'plugin is loaded, presence of Last/Insert/Bottom/Foot data, edition, license status, '
         'document name, and object/layer counts.'),
    ]

    for cmd_name, desc in export_commands:
        doc.add_heading(cmd_name, level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # =====================================================================
    # 17. DATA MODELS & PARAMETERS
    # =====================================================================
    doc.add_heading('17. Data Models & Parameters', level=1)
    doc.add_paragraph(
        'The plugin uses four primary data models to represent footwear components. Each model '
        'stores hundreds of parameters that define the 3D geometry.'
    )

    doc.add_heading('Last Model', level=2)
    doc.add_paragraph(
        'The Last class stores all parameters for the shoe last including: stick length, ball width, '
        'heel width, ball girth, instep girth, waist girth, heel height, toe spring, ball break angle, '
        'cross-section profiles (A1/A2/A3 sections with lateral and medial measurements), girth curve '
        'data, and geometry IDs for all generated objects.'
    )

    doc.add_heading('Insert Model', level=2)
    doc.add_paragraph(
        'The Insert class manages insole/insert parameters: base thickness, top cover, bottom cover, '
        'arch height, heel cup depth, posting angles (medial/lateral), material type, footbed depth, '
        'and contour profiles.'
    )

    doc.add_heading('Bottom Model', level=2)
    doc.add_paragraph(
        'The Bottom class stores outsole, midsole, insole board, shank board, and welt thickness '
        'values, plus sole profile type and construction method.'
    )

    doc.add_heading('Foot Model', level=2)
    doc.add_paragraph(
        'The Foot class stores imported foot scan data, plantar metrics (length, widths, arch height), '
        'and references to scan mesh objects in the document.'
    )

    doc.add_heading('Material Thicknesses', level=2)
    doc.add_paragraph(
        'The MaterialThicknesses class manages thickness values for all material layers:'
    )
    thickness_items = [
        'Insole: base, top cover, bottom cover, posting (medial/lateral)',
        'Bottom: outsole, midsole, insole board, shank, welt',
        'Computed totals: total insole thickness, total bottom thickness, total build height',
    ]
    for item in thickness_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Document Settings', level=2)
    doc.add_paragraph(
        'DocumentSettings provides a key-value store for per-document configuration. All settings '
        'are serialized to JSON and stored in the .3dm file\'s user text strings. Settings include '
        'last size, size system, material choices, fit parameters, cross-section positions, and more.'
    )

    doc.add_page_break()

    # =====================================================================
    # 18. LAYER STRUCTURE
    # =====================================================================
    doc.add_heading('18. Layer Structure', level=1)
    doc.add_paragraph(
        'The plugin creates and manages a hierarchical layer structure in the Rhino document. '
        'All plugin layers are children of the top-level "SLM" (Shoe Last Maker) parent layer.'
    )

    layers = [
        ('SLM', 'Parent layer for all plugin geometry'),
        ('SLM::Last', 'Shoe last surface and construction geometry'),
        ('SLM::Insert', 'Insole/insert geometry'),
        ('SLM::Bottom', 'Bottom components (outsole, midsole, shank, etc.)'),
        ('SLM::Foot', 'Imported foot scan data'),
        ('SLM::Orthotic', 'Orthotic device geometry'),
        ('SLM::Sandal', 'Sandal-specific components'),
        ('SLM::Measurements', 'Measurement annotation lines and curves'),
        ('SLM::ClippingPlanes', 'Cross-section clipping planes'),
        ('SLM::Flattened', 'Flattened 2D patterns'),
        ('SLM::PrintPrep', 'Print-ready geometry'),
    ]

    table = doc.add_table(rows=len(layers) + 1, cols=2)
    table.style = 'Light Grid Accent 1'
    cells = table.rows[0].cells
    cells[0].text = 'Layer'
    cells[1].text = 'Purpose'
    for cell in cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for i, (name, purpose) in enumerate(layers):
        cells = table.rows[i + 1].cells
        cells[0].text = name
        cells[1].text = purpose

    doc.add_page_break()

    # =====================================================================
    # 19. DOCUMENT PERSISTENCE
    # =====================================================================
    doc.add_heading('19. Document Persistence', level=1)
    doc.add_paragraph(
        'All plugin data is automatically saved inside the .3dm file using Rhino\'s document '
        'user text (doc.Strings) system. This means no external files are needed - everything '
        'travels with the Rhino document.'
    )

    doc.add_heading('What is Saved', level=2)
    saved_items = [
        'Geometry references - Base64-encoded serialized geometry keyed by name',
        'Document settings - JSON object with all design parameters',
        'Material thicknesses - JSON object with all thickness values',
        'Plugin version - Version stamp for compatibility checking',
    ]
    for item in saved_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Legacy Format Support', level=2)
    doc.add_paragraph(
        'The plugin supports reading both the current JSON format and the legacy pipe-delimited '
        'format (key1|value1||key2|value2||...) used by older .NET plugin versions. This ensures '
        'backward compatibility with files created by the original 3DShoemaker plugin.'
    )

    doc.add_heading('Automatic Save/Load', level=2)
    doc.add_paragraph(
        'The plugin hooks into Rhino\'s document events:'
    )
    events = [
        'BeginSaveDocument - Writes all plugin data to document user text',
        'EndOpenDocument - Reads plugin data and sets up layers',
        'CloseDocument - Cleans up in-memory caches',
        'NewDocument - Sets up default layers, rendering, and viewport',
    ]
    for event in events:
        doc.add_paragraph(event, style='List Bullet')

    doc.add_page_break()

    # =====================================================================
    # 20. TROUBLESHOOTING
    # =====================================================================
    doc.add_heading('20. Troubleshooting', level=1)

    issues = [
        ('Plugin does not load', 'Verify that the plugin files are in the correct PythonPlugIns '
         'directory for Rhino 8. Check the Rhino command line for error messages during startup. '
         'Ensure the manifest.yml file exists alongside __init__.py.'),
        ('License activation fails', 'Check your internet connection. Verify the license key format '
         '(alphanumeric with dashes). If the server is unreachable, try again later. Contact '
         'support at ShoeLastMaker.com if the issue persists.'),
        ('Commands not found', 'Type the exact command name (case-sensitive) in the Rhino command '
         'line. Verify the plugin loaded successfully by checking for "[3DShoemaker]" messages '
         'in the command history. Run the Testing command for diagnostics.'),
        ('Geometry creation fails', 'Ensure a last has been created or imported first. Most '
         'component commands require an active last. Check that the current document has the '
         'SLM layer hierarchy set up (run Establish if needed).'),
        ('Grading produces incorrect results', 'Verify the size system setting matches your '
         'intended system. Check that the current size is set correctly in document settings. '
         'The grading origin should be at the heel center for proper scaling.'),
        ('Import fails for foot scans', 'Ensure the file format is supported. For point clouds, '
         'verify the file uses comma, space, or tab-separated XYZ values. For mesh files, check '
         'that the file is not corrupted.'),
        ('Morph produces unexpected results', 'Check that source and target point counts match '
         'for P2P morphing. For M2M, ensure the source and target meshes have similar topology. '
         'Adjust the tolerance value if deformation is too smooth or too sharp.'),
    ]

    for title_text, solution in issues:
        doc.add_heading(title_text, level=2)
        doc.add_paragraph(solution)

    doc.add_page_break()

    # =====================================================================
    # 21. TERMS AND CONDITIONS
    # =====================================================================
    doc.add_heading('21. Terms and Conditions', level=1)
    doc.add_paragraph(
        'The use of the 3DShoemaker plugin is governed by the Terms and Conditions available at '
        'https://ShoeLastMaker.com and included in the Terms.txt file distributed with the plugin.'
    )
    doc.add_paragraph('Key points:')
    terms_points = [
        'A valid license key is required, purchasable at 3DShoemaker.com.',
        'Personal Edition may not be used for commercial purposes.',
        'Each license is tied to a single machine at a time.',
        'The software may not be copied, modified, or reverse-engineered.',
        'Customer outputs (models, designs, products) are owned by the customer.',
        'The software is provided "as is" without warranty.',
        'The software does not provide medical or clinical advice.',
        'An active internet connection is required to operate.',
    ]
    for point in terms_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_paragraph('')
    doc.add_paragraph('')

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('3DShoemaker Plugin for Rhino 8 - User Manual v8.4.0.8')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run = footer.add_run('\nhttps://ShoeLastMaker.com')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x66, 0xCC)

    # Save
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        '3DShoemaker_Rhino8_User_Manual.docx'
    )
    doc.save(output_path)
    print(f'Manual saved to: {output_path}')
    return output_path


if __name__ == '__main__':
    create_manual()
